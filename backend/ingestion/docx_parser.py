"""
ingestion/docx_parser.py - DOCX text extraction for Kizlly.

Uses python-docx to extract paragraphs and table contents, returning the
same dict structure as pdf_parser for uniform downstream processing.
"""

from __future__ import annotations

import logging
import os
from typing import Dict

logger = logging.getLogger(__name__)

try:
    from docx import Document  # python-docx
except ImportError:
    Document = None  # type: ignore[misc,assignment]
    logger.warning("python-docx is not installed; DOCX parsing unavailable.")


def parse_docx(file_path: str) -> Dict:
    """Parse a DOCX file and return structured text with metadata.

    Parameters
    ----------
    file_path : str
        Absolute or relative path to the ``.docx`` file.

    Returns
    -------
    dict
        ``{"text": str, "pages": int, "metadata": dict}``

        *pages* is set to ``1`` because DOCX files do not encode physical
        page boundaries; the entire content is treated as a single logical
        page with a ``[PAGE 1]`` marker for consistency with the PDF parser.

    Raises
    ------
    FileNotFoundError
        If *file_path* does not exist.
    RuntimeError
        If python-docx is not installed or extraction fails.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    try:
        doc = Document(file_path)
    except Exception as exc:
        raise RuntimeError(f"Failed to open DOCX file {file_path}: {exc}") from exc

    # ------------------------------------------------------------------
    # Extract metadata from core properties
    # ------------------------------------------------------------------
    metadata: dict = {}
    try:
        cp = doc.core_properties
        for attr in ("author", "title", "subject", "created", "modified"):
            val = getattr(cp, attr, None)
            if val:
                metadata[attr] = str(val)
    except Exception:
        pass  # metadata extraction is best-effort

    # ------------------------------------------------------------------
    # Extract paragraphs
    # ------------------------------------------------------------------
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # ------------------------------------------------------------------
    # Extract tables
    # ------------------------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "[PAGE 1]\n" + "\n\n".join(parts) if parts else ""
    page_count = 1 if parts else 0

    logger.info(
        "Parsed DOCX %s: %d paragraphs/rows, %d characters.",
        file_path,
        len(parts),
        len(full_text),
    )

    return {
        "text": full_text,
        "pages": page_count,
        "metadata": metadata,
    }
